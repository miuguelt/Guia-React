"""
Generador de Guia de Aprendizaje - SENA ADSO
Guia React - version DOCX
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CODIGO_DIR = os.path.join(BASE_DIR, "recursos", "codigo-ejemplo")


def leer_archivo(ruta_relativa):
    ruta = os.path.join(CODIGO_DIR, ruta_relativa)
    with open(ruta, "r", encoding="utf-8") as f:
        return f.read()


def configurar_estilos(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for nivel in range(1, 4):
        h = doc.styles[f"Heading {nivel}"]
        h.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        h.font.name = "Calibri"

    try:
        code_style = doc.styles.add_style("Codigo", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        code_style = doc.styles["Codigo"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0


def agregar_portada(doc):
    for _ in range(4):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SERVICIO NACIONAL DE APRENDIZAJE - SENA")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tecnologo en Analisis y Desarrollo de Software (ADSO)")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GUIA DE APRENDIZAJE")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Desarrollo de Interfaces con React\nVite, Hooks y JavaScript ES6+")
    run.font.size = Pt(14)

    for _ in range(4):
        doc.add_paragraph("")
    tabla = doc.add_table(rows=5, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    datos = [("Version", "1.0"), ("Fecha", "Junio 2026"), ("Autor", "Instructor ADSO - SENA"), ("Duracion", "40 horas"), ("Modalidad", "Presencial / Virtual")]
    for i, (k, v) in enumerate(datos):
        run_k = tabla.rows[i].cells[0].paragraphs[0].add_run(k)
        run_k.font.bold = True
        tabla.rows[i].cells[1].paragraphs[0].add_run(v)
    doc.add_page_break()


def agregar_bloque_codigo(doc, codigo):
    for linea in codigo.split("\n"):
        p = doc.add_paragraph()
        p.style = doc.styles["Codigo"]
        p.paragraph_format.left_indent = Cm(1)
        shading = p._element.get_or_add_pPr()
        shade = shading.makeelement(qn("w:shd"), {qn("w:fill"): "F5F5F5", qn("w:val"): "clear"})
        shading.append(shade)
        run = p.add_run(linea if linea else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def agregar_bloque_codigo_con_ruta(doc, ruta, codigo):
    p = doc.add_paragraph()
    run = p.add_run(f"📄 {ruta}")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    agregar_bloque_codigo(doc, codigo)


def generar_documento():
    doc = Document()
    configurar_estilos(doc)

    print("Generando portada...")
    agregar_portada(doc)

    print("Generando contenido...")
    doc.add_heading("1. IDENTIFICACION", level=1)
    doc.add_paragraph("Guia de aprendizaje para desarrollo con React, Vite y JavaScript ES6+.")

    doc.add_heading("2. PRESENTACION", level=1)
    doc.add_paragraph("React es la libreria de UI mas popular del ecosistema JavaScript. Esta guia te llevara desde cero hasta construir aplicaciones web modernas con componentes, hooks y routing.")

    doc.add_heading("3. ACTIVIDADES DE APRENDIZAJE", level=1)
    doc.add_heading("3.1 Reflexion Inicial", level=2)
    doc.add_paragraph("Antes de React, manipular el DOM era complejo y propenso a errores. React introdujo el Virtual DOM y los componentes declarativos.")

    doc.add_heading("3.2 Configuracion del Entorno", level=2)
    agregar_bloque_codigo_con_ruta(doc, "powershell", "npm create vite@latest mi-app -- --template react\ncd mi-app\nnpm install\nnpm run dev")

    doc.add_heading("3.3 Componentes y Props", level=2)
    agregar_bloque_codigo_con_ruta(doc, "src/components/Saludo.jsx", leer_archivo("src/components/Saludo.jsx"))

    doc.add_heading("3.4 Estado con useState", level=2)
    agregar_bloque_codigo_con_ruta(doc, "src/components/Contador.jsx", leer_archivo("src/components/Contador.jsx"))

    doc.add_heading("3.5 Efectos con useEffect", level=2)
    agregar_bloque_codigo_con_ruta(doc, "src/components/Reloj.jsx", leer_archivo("src/components/Reloj.jsx"))

    doc.add_heading("3.6 Custom Hooks", level=2)
    agregar_bloque_codigo_con_ruta(doc, "src/hooks/useFetch.js", leer_archivo("src/hooks/useFetch.js"))

    doc.add_heading("4. EVIDENCIAS", level=1)
    doc.add_paragraph("Para demostrar los resultados de aprendizaje, el aprendiz debe construir:")
    items = ["App de Tareas con useState y localStorage", "Frontend Fincas y Cultivos conectando a Spring Boot", "Dashboard con Custom Hooks y Context"]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    ruta_salida = os.path.join(BASE_DIR, "Guia_Aprendizaje_React.docx")
    doc.save(ruta_salida)
    print(f"Documento generado: {ruta_salida}")


if __name__ == "__main__":
    generar_documento()
